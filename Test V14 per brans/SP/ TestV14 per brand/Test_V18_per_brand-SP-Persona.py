import streamlit as st
from PIL import Image
import requests
from io import BytesIO, BytesIO as io_bytes
from groq import Groq
from datetime import datetime
import re
import fitz  # PyMuPDF
import pdfplumber
import base64
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# --- Optional Word download ---
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ python-docx not installed. Word download unavailable.")

# --- Initialize Groq client with hardcoded API key ---
GROQ_API_KEY = "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"
client = Groq(api_key=GROQ_API_KEY)

# --- Session state ---
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# --- Language ---
language = st.radio("Select Language / اختر اللغة", options=["English", "العربية"])

# --- GSK Logo ---
logo_local_path = "images/gsk_logo.png"
logo_fallback_url = "https://www.tungsten-network.com/wp-content/uploads/2020/05/GSK_Logo_Full_Colour_RGB.png"
col1, col2 = st.columns([1,5])
with col1:
    try:
        logo_img = Image.open(logo_local_path)
        st.image(logo_img, width=120)
    except:
        st.image(logo_fallback_url, width=120)
with col2:
    st.title("🧠 AI Sales Call Assistant with Evidence & Figures")

# --- Brand & product data ---
gsk_brands = {
    "Shingrix": "Test V14 per brans/SP/ TestV14 per brand/Shingrix.pdf",
}
gsk_brands_images = {
    "Shingrix": "https://www.oma-apteekki.fi/WebRoot/NA/Shops/na/67D6/48DA/D0B0/D959/ECAF/0A3C/0E02/D573/3ad67c4e-e1fb-4476-a8a0-873423d8db42_3Dimage.png",
}

# --- Filters & options ---
race_segments = [
    "R – Reach: Did not start to prescribe yet and Don't believe that vaccination is his responsibility.",
    "A – Acquisition: Prescribe to patient who initiate discussion about the vaccine but Convinced about Shingrix data.",
    "C – Conversion: Proactively initiate discussion with specific patient profile but For other patient profiles he is not prescribing yet.",
    "E – Engagement: Proactively prescribe to different patient profiles"
]
doctor_barriers = [
    "HCP does not consider HZ as risk",
    "No time to discuss preventive measures",
    "Cost considerations",
    "Not convinced HZ Vx effective",
    "Accessibility issues"
]
objectives = ["Awareness", "Adoption", "Retention"]
specialties = ["GP", "Cardiologist", "Dermatologist", "Endocrinologist", "Pulmonologist"]
personas = [
    "Uncommitted Vaccinator",
    "Reluctant Efficiency",
    "Patient Influenced",
    "Committed Vaccinator"
]
gsk_approaches = [
    "Use data-driven evidence",
    "Focus on patient outcomes",
    "Leverage storytelling techniques"
]
sales_call_flow = ["Prepare", "Engage", "Create Opportunities", "Influence", "Drive Impact", "Post Call Analysis"]

# --- Sidebar filters ---
st.sidebar.header("Filters & Options")
brand = st.sidebar.selectbox("Select Brand / اختر العلامة التجارية", options=list(gsk_brands.keys()))
segment = st.sidebar.selectbox("Select RACE Segment / اختر شريحة RACE", race_segments)
barrier = st.sidebar.multiselect("Select Doctor Barrier / اختر حاجز الطبيب", options=doctor_barriers, default=[])
objective = st.sidebar.selectbox("Select Objective / اختر الهدف", options=objectives)
specialty = st.sidebar.selectbox("Select Doctor Specialty / اختر تخصص الطبيب", options=specialties)
persona = st.sidebar.selectbox("Select HCP Persona / اختر شخصية الطبيب", options=personas)
response_length = st.sidebar.selectbox("Response Length / اختر طول الرد", ["Short", "Medium", "Long"])
response_tone = st.sidebar.selectbox("Response Tone / اختر نبرة الرد", ["Formal", "Casual", "Friendly", "Persuasive"])
interface_mode = st.sidebar.radio("Interface Mode / اختر واجهة", ["Chatbot", "Card Dashboard", "Flow Visualization"])

# --- Load Shingrix PDF (text + figures) ---
pdf_text, pdf_figures = "", []
pdf_path = gsk_brands[brand]

try:
    # Extract text
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            pdf_text += page.extract_text() or ""
    pdf_text_truncated = pdf_text[:2000]

    # Extract figures + captions
    doc = fitz.open(pdf_path)
    for page in doc:
        blocks = page.get_text("blocks")
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]

            rect = fitz.Rect(page.get_image_bbox(img))
            caption_text = ""
            for block in blocks:
                bx, by, ex, ey, _, text, *_ = block
                block_rect = fitz.Rect(bx, by, ex, ey)
                if rect.intersects(block_rect) or abs(block_rect.y0 - rect.y1) < 50:
                    if isinstance(text, str):
                        caption_text += " " + text.strip()
            caption_text = caption_text.strip() or f"Figure {len(pdf_figures)+1}"
            pdf_figures.append({"image": img_bytes, "caption": caption_text})

    st.success("✅ Shingrix leaflet loaded with text and figures.")
except Exception as e:
    st.warning(f"⚠️ Could not process Shingrix PDF: {e}")

# --- Display brand image ---
image_path = gsk_brands_images.get(brand)
try:
    if image_path.startswith("http"):
        response = requests.get(image_path)
        img = Image.open(BytesIO(response.content))
    else:
        img = Image.open(image_path)
    st.image(img, width=200)
except:
    st.warning(f"⚠️ Could not load image for {brand}. Using placeholder.")
    st.image("https://via.placeholder.com/200x100.png?text=No+Image", width=200)

# --- Clear chat ---
if st.button("🗑️ Clear Chat / مسح المحادثة"):
    st.session_state.chat_history = []

# --- Chat display ---
st.subheader("💬 Chatbot Interface")
chat_placeholder = st.empty()

def display_chat(selected_figures=None):
    chat_html = ""
    for msg in st.session_state.chat_history:
        time = msg.get("time", "")
        content = msg["content"].replace('\n', '<br>')

        # Embed only selected figures
        if selected_figures:
            for f in selected_figures:
                caption = f["caption"]
                img_bytes = f["image"]
                img_html = f'<img src="data:image/png;base64,{base64.b64encode(img_bytes).decode()}" width="400"/>'
                content = content.replace(caption, f"{caption}<br>{img_html}")

        if msg["role"] == "user":
            chat_html += f"<div style='text-align:right; background:#dcf8c6; padding:10px; border-radius:15px 15px 0px 15px; margin:5px; display:inline-block; max-width:80%;'>{content}<span style='font-size:10px; color:gray;'><br>{time}</span></div>"
        else:
            chat_html += f"<div style='text-align:left; background:#f0f2f6; padding:10px; border-radius:15px 15px 15px 0px; margin:5px; display:inline-block; max-width:80%;'>{content}<span style='font-size:10px; color:gray;'><br>{time}</span></div>"
    chat_placeholder.markdown(chat_html, unsafe_allow_html=True)

display_chat()

# --- Chat input ---
with st.form("chat_form", clear_on_submit=True):
    user_input = st.text_input("Type your message...", key="user_input_box")
    submitted = st.form_submit_button("➤")

if submitted and user_input.strip():
    st.session_state.chat_history.append({"role": "user", "content": user_input, "time": datetime.now().strftime("%H:%M")})

    # --- Select most relevant figures ---
    def select_relevant_figures(user_query, figures, top_k=2):
        if not figures:
            return []
        captions = [f["caption"] for f in figures]
        corpus = captions + [user_query]
        vectorizer = TfidfVectorizer().fit_transform(corpus)
        similarity = cosine_similarity(vectorizer[-1:], vectorizer[:-1])
        top_indices = similarity[0].argsort()[::-1][:top_k]
        return [figures[i] for i in top_indices]

    selected_figures = select_relevant_figures(user_input, pdf_figures, top_k=2)

    approaches_str = "\n".join(gsk_approaches)
    flow_str = " → ".join(sales_call_flow)
    figure_texts = "\n".join([f"{i+1}. {f['caption']}" for i, f in enumerate(selected_figures)])

    prompt = f"""
Language: {language}
User input: {user_input}
RACE Segment: {segment}
Doctor Barrier: {', '.join(barrier) if barrier else 'None'}
Objective: {objective}
Brand: {brand}
Doctor Specialty: {specialty}
HCP Persona: {persona}

Approved Sales Approaches:
{approaches_str}

Sales Call Flow Steps:
{flow_str}

Use APACT technique for objections.
Response Length: {response_length}
Response Tone: {response_tone}

Leaflet Evidence:
{pdf_text_truncated}

Include relevant figures from the leaflet in your response when appropriate. Reference figures by their captions.

FIGURES:
{figure_texts}
"""

    # --- Call Groq API ---
    response = client.chat.completions.create(
        model="llama-3.1-70b-versatile",
        messages=[
            {"role": "system", "content": f"You are a helpful sales assistant chatbot that responds in {language}."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7
    )

    ai_output = response.choices[0].message.content
    st.session_state.chat_history.append({"role": "ai", "content": ai_output, "time": datetime.now().strftime("%H:%M")})
    display_chat(selected_figures=selected_figures)

# --- Word download ---
if DOCX_AVAILABLE and st.session_state.chat_history:
    latest_ai = [msg["content"] for msg in st.session_state.chat_history if msg["role"] == "ai"]
    if latest_ai:
        doc = Document()
        doc.add_heading("AI Sales Call Response", 0)
        doc.add_paragraph(latest_ai[-1])
        word_buffer = io_bytes()
        doc.save(word_buffer)
        st.download_button("📥 Download as Word (.docx)", word_buffer.getvalue(), file_name="AI_Response.docx")

# --- Brand leaflet link ---
st.markdown(f"[Brand Leaflet - {brand}]({gsk_brands[brand]})")
