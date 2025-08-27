import streamlit as st
from PIL import Image
import requests
from io import BytesIO, BytesIO as io_bytes
import groq
from groq import Groq
from datetime import datetime

# Optional: Word download
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("⚠️ python-docx not installed. Word download unavailable.")

# PDF dependencies
try:
    import PyPDF2
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    st.warning("⚠️ PyPDF2 or PyMuPDF not installed. PDF features disabled.")

# --- Groq client ---
client = Groq(api_key="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")  # 🔑 Replace with your Groq API key

# --- Session state ---
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "pdf_text" not in st.session_state:
    st.session_state.pdf_text = ""
if "pdf_images" not in st.session_state:
    st.session_state.pdf_images = []
if "pdf_visual_text" not in st.session_state:
    st.session_state.pdf_visual_text = ""
if "pdf_failed" not in st.session_state:
    st.session_state.pdf_failed = False

# --- Language ---
language = st.radio("Select Language / اختر اللغة", options=["English", "العربية"])

# --- GSK Logo ---
logo_url = "https://www.tungsten-network.com/wp-content/uploads/2020/05/GSK_Logo_Full_Colour_RGB.png"
st.image(logo_url, width=150)
st.title("🧠 AI Sales Call Assistant")

# --- Brands + GitHub PDF URLs (raw links) ---
brand_pdfs = {
    "Shingrix": "https://raw.githubusercontent.com/karimgsk6-debug/New-AI-sales-call-assistant/main/TestV14_per_brans/SP/TestV14_per_brand/Shingrix.pdf",
    "Trelegy": "https://raw.githubusercontent.com/karimgsk6-debug/New-AI-sales-call-assistant/main/TestV14_per_brans/SP/TestV14_per_brand/Trelegy.pdf",
    "Zejula": "https://raw.githubusercontent.com/karimgsk6-debug/New-AI-sales-call-assistant/main/TestV14_per_brans/SP/TestV14_per_brand/Zejula.pdf"
}

# --- Filters ---
race_segments = ["R – Reach", "A – Acquisition", "C – Conversion", "E – Engagement"]
doctor_barriers = ["HCP does not consider HZ as risk", "No time", "Cost", "Not convinced effective", "Accessibility issues"]
objectives = ["Awareness", "Adoption", "Retention"]
specialties = ["GP", "Cardiologist", "Dermatologist", "Endocrinologist", "Pulmonologist"]
personas = ["Uncommitted Vaccinator", "Reluctant Efficiency", "Patient Influenced", "Committed Vaccinator"]
gsk_approaches = ["Use data-driven evidence", "Focus on patient outcomes", "Leverage storytelling techniques"]
sales_call_flow = ["Prepare", "Engage", "Create Opportunities", "Influence", "Drive Impact", "Post Call Analysis"]

# --- Sidebar filters ---
st.sidebar.header("Filters & Options")
brand = st.sidebar.selectbox("Select Brand", options=list(brand_pdfs.keys()))
segment = st.sidebar.selectbox("Select RACE Segment", race_segments)
barrier = st.sidebar.multiselect("Select Doctor Barrier", doctor_barriers, default=[])
objective = st.sidebar.selectbox("Select Objective", objectives)
specialty = st.sidebar.selectbox("Select Doctor Specialty", specialties)
persona = st.sidebar.selectbox("Select HCP Persona", personas)
response_length = st.sidebar.selectbox("Response Length", ["Short", "Medium", "Long"])
response_tone = st.sidebar.selectbox("Response Tone", ["Formal", "Casual", "Friendly", "Persuasive"])

# --- Fetch PDF content automatically ---
if PDF_AVAILABLE:
    try:
        pdf_url = brand_pdfs[brand]
        r = requests.get(pdf_url)
        r.raise_for_status()

        # Check if it's really a PDF
        if not r.content.startswith(b"%PDF"):
            raise ValueError("File is not a valid PDF (GitHub may be serving HTML instead).")

        pdf_file = BytesIO(r.content)

        # Extract text
        reader = PyPDF2.PdfReader(pdf_file)
        pdf_text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                pdf_text += extracted + "\n"
        st.session_state.pdf_text = pdf_text[:6000]

        # Extract visuals
        pdf_file.seek(0)
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        images = []
        image_descriptions = []
        for page_num in range(len(doc)):
            for img_index, img in enumerate(doc[page_num].get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_data = base_image["image"]
                pil_img = Image.open(BytesIO(image_data))
                images.append(pil_img)

                # --- OPTIONAL: OCR on visuals ---
                try:
                    import pytesseract
                    text_in_image = pytesseract.image_to_string(pil_img)
                    if text_in_image.strip():
                        image_descriptions.append(text_in_image.strip())
                except ImportError:
                    image_descriptions.append(f"Visual {len(images)} extracted (OCR not available).")

            if len(images) >= 3:  # limit to first 3 visuals
                break

        st.session_state.pdf_images = images
        st.session_state.pdf_visual_text = "\n".join(image_descriptions)
        st.session_state.pdf_failed = False

    except Exception as e:
        st.warning(f"⚠️ Could not fetch or parse PDF: {e}")
        st.session_state.pdf_text = ""
        st.session_state.pdf_images = []
        st.session_state.pdf_visual_text = ""
        st.session_state.pdf_failed = True
else:
    st.session_state.pdf_text = ""
    st.session_state.pdf_images = []
    st.session_state.pdf_visual_text = ""
    st.session_state.pdf_failed = True

# --- Chat input ---
st.subheader("💬 Chatbot Interface")
with st.form("chat_form", clear_on_submit=True):
    user_input = st.text_input("Type your message...")
    submitted = st.form_submit_button("➤")

if submitted and user_input.strip():
    st.session_state.chat_history.append({"role": "user", "content": user_input, "time": datetime.now().strftime("%H:%M")})

    pdf_context = st.session_state.pdf_text if st.session_state.pdf_text else "No leaflet content available."
    visual_context = st.session_state.pdf_visual_text if st.session_state.pdf_visual_text else "No visual info available."

    prompt = f"""
Language: {language}
User Input: {user_input}
Brand: {brand}
Objective: {objective}
RACE Segment: {segment}
Doctor Barriers: {', '.join(barrier) if barrier else 'None'}
Specialty: {specialty}
Persona: {persona}

--- Official Leaflet Info (Text) ---
{pdf_context}

--- Official Leaflet Visuals (OCR / extracted text) ---
{visual_context}

--- Instructions ---
Use APACT method (Acknowledge → Probing → Answer → Confirm → Transition).
Provide actionable sales suggestions aligned with BOTH the text and visuals of the leaflet.
Adjust tone: {response_tone}, Length: {response_length}.
"""

    response = client.chat.completions.create(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[
            {"role": "system", "content": f"You are a helpful sales assistant in {language}."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7
    )

    ai_output = response.choices[0].message.content
    st.session_state.chat_history.append({"role": "ai", "content": ai_output, "time": datetime.now().strftime("%H:%M")})

# --- Display chat ---
for msg in st.session_state.chat_history:
    if msg["role"] == "user":
        st.markdown(f"🧑‍💼 **You:** {msg['content']}")
    else:
        st.markdown(f"🤖 **AI:** {msg['content']}")

# --- Display visuals ---
if st.session_state.pdf_images:
    st.subheader("📊 Extracted Visuals from Leaflet")
    for img in st.session_state.pdf_images:
        st.image(img, use_container_width=True)

# --- Fallback PDF viewer if parsing failed ---
if st.session_state.pdf_failed:
    st.subheader("📖 Leaflet Viewer (Fallback)")
    pdf_url = brand_pdfs[brand]
    st.markdown(
        f'<iframe src="{pdf_url}" width="100%" height="600px" style="border: none;"></iframe>',
        unsafe_allow_html=True
    )

# --- Word download ---
if DOCX_AVAILABLE and st.session_state.chat_history:
    latest_ai = [msg["content"] for msg in st.session_state.chat_history if msg["role"] == "ai"]
    if latest_ai:
        doc = Document()
        doc.add_heading("AI Sales Call Response", 0)
        doc.add_paragraph(latest_ai[-1])
        word_buffer = io_bytes()
        doc.save(word_buffer)
        st.download_button("📥 Download as Word (.docx)", word_buffer.getvalue(), file_name="AI_Response.docx")
